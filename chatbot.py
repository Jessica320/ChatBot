import os
import requests
import warnings
import chromadb
from requests.exceptions import SSLError, RequestException
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain_core.runnables import Runnable
from langchain_core.retrievers import BaseRetriever
from langchain_core.documents import Document
from typing import List, Any, Dict
from chromadb.utils import embedding_functions

# 引入 Guardrails 相關功能
import ner_guardrails

# 隱藏 InsecureRequestWarning
warnings.filterwarnings("ignore", category=requests.packages.urllib3.exceptions.InsecureRequestWarning)

# API 參數
API_URL = 'https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent'
MAX_NEW_TOKENS = 800
TEMPERATURE = 0.8
TOP_K = 60
TOP_P = 0.9

# 嵌入模型和 Chroma 參數
MODEL_PATH = "./paraphrase-multilingual-MiniLM-L12-v2"
PDF_DIR = "./KM_pool"
CHROMA_PATH = "./chroma_db"

API_KEY_FILE = os.path.join(os.path.dirname(__file__), "apikey.txt")

# 從檔案讀取 API 金鑰
def load_api_key(api_key_file):
    with open(api_key_file, 'r', encoding='utf-8') as f:
        api_key = f.read().strip()
        if not api_key:
            print(f"錯誤：檔案 {api_key_file} 為空")
            return None
        return api_key

# 自定義 Gemini LLM 類
class GeminiAPI(Runnable):
    def __init__(self, api_key: str, api_url: str, max_new_tokens: int, temperature: float, top_k: int = None, top_p: int = None):
        super().__init__()
        self.api_key = api_key
        self.api_url = api_url
        self.max_new_tokens = max_new_tokens
        self.temperature = temperature
        self.top_k = top_k
        self.top_p = top_p

    def invoke(self, input: Any, config: Dict = None, **kwargs) -> str:
        prompt = str(input)
        context_str = ""
        history_str = ""

        if "context" in kwargs and kwargs["context"]:
            # 對檢索到的上下文進行脫敏處理
            desensitized_context = []
            for doc in kwargs["context"]:
                entities = ner_guardrails.extract_entities_with_regex(doc.page_content)
                desensitized_content = ner_guardrails.desensitize_text_with_entities(doc.page_content, entities)
                desensitized_context.append(Document(page_content=desensitized_content, metadata=doc.metadata))
            context = "\n".join([doc.page_content for doc in desensitized_context])
            context_str = f"上下文：\n{context}\n\n"

        if "chat_history" in kwargs and kwargs["chat_history"]:
            history_str = "之前的對話：\n"
            for msg in kwargs["chat_history"]:
                history_str += f"{msg.type}: {msg.content}\n"
            history_str += "\n"

        # 更詳細地指示模型參考歷史，並在回答中考慮先前的問題和答案
        full_prompt = f"請使用中文回答以下問題，並仔細參考之前的對話歷史。在你的回答中，考慮先前人類提出的問題以及你給出的答案，以確保回答的連貫性。不要使用英文。\n\n{history_str}{context_str}人類：{prompt}"

        headers = {"Content-Type": "application/json", "x-goog-api-key": self.api_key}
        payload = {
            "contents": [{"parts": [{"text": full_prompt}]}],
            "generationConfig": {
                "temperature": self.temperature,
                "maxOutputTokens": self.max_new_tokens,
                "topK": self.top_k,
                "topP": self.top_p
            }
        }

        try:
            response = requests.post(self.api_url, headers=headers, json=payload, verify=False)
            response.raise_for_status()
            result = response.json()
            generated_text = result["candidates"][0]["content"]["parts"][0]["text"]
            tokens = generated_text.split()
            if len(tokens) > self.max_new_tokens:
                generated_text = " ".join(tokens[:self.max_new_tokens]) + "..."
            return generated_text
        except (SSLError, RequestException, KeyError, IndexError) as e:
            print(f"API 請求或解析失敗: {e}")
            return "無法生成回答"

    def _get_input_schema(self, config=None):
        return str

    def _get_output_schema(self, config=None):
        return str

# 初始化嵌入函數
def init_embedding_function():
    return embedding_functions.SentenceTransformerEmbeddingFunction(model_name=MODEL_PATH)

# 載入或建立 Chroma 資料庫
def setup_vectorstore():
    if not os.path.exists(MODEL_PATH):
        raise FileNotFoundError(f"嵌入模型路徑 {MODEL_PATH} 不存在")
    if not os.path.exists(PDF_DIR):
        raise FileNotFoundError(f"檔案目錄 {PDF_DIR} 不存在")

    client = chromadb.PersistentClient(path=CHROMA_PATH)
    collection_name = "pdf_docx_collection"
    embedding_function = init_embedding_function()

    try:
        collection = client.get_collection(name=collection_name, embedding_function=embedding_function)
        print(f"成功載入現有集合：{collection.name}，來自 {CHROMA_PATH}")
        return collection
    except Exception as e:
        print(f"集合 {collection_name} 不存在，將創建新集合。")
        collection = client.create_collection(name=collection_name, embedding_function=embedding_function)
        # 在這裡直接處理檔案並添加到集合
        all_files = [os.path.join(PDF_DIR, f) for f in os.listdir(PDF_DIR) if f.lower().endswith((".pdf", ".docx"))]
        if not all_files:
            print(f"警告：目錄 {PDF_DIR} 中找不到 PDF 或 DOCX 檔案")
            return collection
        all_documents = []
        all_metadatas = []
        all_ids = []
        for file_path in all_files:
            file_name = os.path.basename(file_path)
            print(f"處理檔案：{file_name}")
            if file_path.lower().endswith(".pdf"):
                from PyPDF2 import PdfReader
                try:
                    reader = PdfReader(file_path)
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text() or ""
                        text = " ".join(text.split())
                        if text:
                            all_documents.append(f"檔案名稱：{file_name}\n頁碼：{page_num}\n內容：{text}")
                            all_metadatas.append({"source": file_name, "page_num": page_num, "file_name": file_name})
                            all_ids.append(f"{file_name}_page{page_num}")
                except Exception as e:
                    print(f"讀取 PDF {file_name} 時發生錯誤：{e}")
            elif file_path.lower().endswith(".docx"):
                from docx import Document as DocxReader
                try:
                    document = DocxReader(file_path)
                    full_text = "\n".join([p.text for p in document.paragraphs])
                    if full_text:
                        all_documents.append(f"檔案名稱：{file_name}\n內容：{full_text}")
                        all_metadatas.append({"source": file_name, "file_name": file_name})
                        all_ids.append(f"{file_name}_docx")
                except Exception as e:
                    print(f"讀取 DOCX {file_name} 時發生錯誤：{e}")
        if all_documents:
            collection.add(documents=all_documents, metadatas=all_metadatas, ids=all_ids)
            print(f"已將 {len(all_files)} 個檔案的內容添加到集合：{collection_name}")
        else:
            print("警告：沒有找到可添加到集合的內容。")
        return collection

# 獲取資料夾中的文件名
def get_available_filenames():
    return [f for f in os.listdir(PDF_DIR) if f.lower().endswith((".pdf", ".docx"))]

# 改進的檢索器
class ChromaRetriever(BaseRetriever):
    collection: Any
    k: int = 3

    def _get_relevant_documents(self, query: str) -> List[Document]:
        target_file_name = None
        query_lower = query.lower().replace("的摘要", "").strip()
        available_files = get_available_filenames()

        # 更積極地尋找問題中包含的文件名（不包含副檔名）
        for file in available_files:
            file_base = os.path.splitext(file)[0].lower()
            if file_base in query_lower:
                target_file_name = file
                break
        # 如果沒有找到不含副檔名的匹配，再嘗試完整文件名匹配
        if not target_file_name:
            for file in available_files:
                if file.lower() == query_lower:
                    target_file_name = file
                    break

        if target_file_name:
            results_with_filename = self.collection.query(
                query_texts=[f"{target_file_name} {query}"],
                n_results=self.k * 2,
                where={"file_name": target_file_name},
                include=["metadatas", "documents"]
            )
            docs = [Document(page_content=doc, metadata=metadata)
                    for doc, metadata in zip(results_with_filename["documents"][0] or [],
                                                results_with_filename["metadatas"][0] or [])]

            if len(docs) < self.k:
                remaining = self.k - len(docs)
                broader_results = self.collection.query(
                    query_texts=[query],
                    n_results=remaining,
                    where={"file_name": {"$ne": target_file_name}},
                    include=["metadatas", "documents"]
                )
                docs.extend(Document(page_content=doc, metadata=metadata)
                            for doc, metadata in zip(broader_results["documents"][0] or [],
                                                        broader_results["metadatas"][0] or []))
            return docs[:self.k]
        else:
            results = self.collection.query(
                query_texts=[query],
                n_results=self.k,
                include=["metadatas", "documents"]
            )
            return [Document(page_content=doc, metadata=metadata)
                    for doc, metadata in zip(results["documents"][0] or [],
                                                results["metadatas"][0] or [])]

def create_rag_chain(api_key: str):
    llm = GeminiAPI(api_key, API_URL, MAX_NEW_TOKENS, TEMPERATURE, TOP_K, TOP_P)
    collection = setup_vectorstore()
    if collection is None:
        print("無法建立 RAG 鏈，因為向量資料庫未成功載入。請檢查是否已運行 embedding.py 建立資料庫。")
        return None
    retriever = ChromaRetriever(collection=collection, k=3)
    memory = ConversationBufferMemory(memory_key="chat_history", input_key="question", output_key="answer", return_messages=True)
    rag_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory, return_source_documents=True, output_key="answer")
    return rag_chain

# 問答函數
def ask_question(rag_chain, question: str, chat_history: List[Any] = None):
    try:
        inputs = {"question": question}
        if chat_history:
            inputs["chat_history"] = chat_history
        result = rag_chain.invoke(inputs)
        answer = result["answer"]
        source_docs = result["source_documents"]
        updated_chat_history = result["chat_history"]

        # 對 LLM 的回答進行脫敏處理
        entities_answer = ner_guardrails.extract_entities_with_regex(answer)
        desensitized_answer = ner_guardrails.desensitize_text_with_entities(answer, entities_answer)

        # 如果沒有檢索到文件且問題包含「摘要」，提供更具體的建議
        if not source_docs and "摘要" in question.lower():
            available_files = get_available_filenames()
            desensitized_answer += f"\n\n⚠️ 無法找到與問題直接相關的文件片段。請嘗試更明確地指定您想查詢的文件名稱，例如：「{available_files[0] if available_files else '文件名' } 的摘要」。可用文件：{', '.join(available_files)}"

        return desensitized_answer, source_docs, updated_chat_history
    except Exception as e:
        print(f"問答過程中發生錯誤：{e}")
        return "無法生成回答，請檢查問題或向量資料庫。", [], chat_history if chat_history else []

# 顯示來源並去重複，現在包含檔案名稱
def process_source_documents(source_docs: List[Any], query: str = "") -> List[str]:
    processed_list = []
    seen_contents = set()
    for doc in source_docs:
        content = doc.page_content.strip()
        source = doc.metadata.get("source", "未知文件")
        if content and content not in seen_contents:
            processed_list.append(f"**來源：{source}**\n\n{content}")
            seen_contents.add(content)
    return processed_list

if __name__ == "__main__":
    api_key = load_api_key(API_KEY_FILE)
    if api_key:
        rag_chain = create_rag_chain(api_key)
        if rag_chain:
            chat_history = []  # 初始化對話歷史
            while True:
                question = input("請輸入您的問題 (輸入 'quit' 結束)：")
                if question.lower() == 'quit':
                    break
                answer, source_docs, chat_history = ask_question(rag_chain, question, chat_history)
                print(f"問題：{question}")
                print(f"答案：{answer}")
                print("\n來源文件：")
                for doc in source_docs:
                    print(f"  - {doc.metadata.get('source', '未知文件')}")
                    # 這裡的 doc.page_content 仍然是原始未脫敏的內容
                    print(f"    {doc.page_content[:100]}...")
                print("\n--- 對話歷史 ---")
                for message in chat_history:
                    print(f"{message.type}: {message.content}")
                print("---")
        else:
            print("RAG 鏈建立失敗，請檢查是否已運行 embedding.py 並建立了向量資料庫。")
    else:
        print("請在 apikey.txt 檔案中提供您的 Gemini API 金鑰。")